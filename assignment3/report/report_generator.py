from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from assignment3.report.experiment_files import *
from assignment3.report.utils import get_descriptive_statistics
import pandas as pd


def _add_page_break(document):
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_image_to_cell(image_path, cell, width_inches, height_inches):
    cell.paragraphs[0]\
        .add_run()\
        .add_picture(image_path, width=Inches(width_inches), height=Inches(height_inches))


def _add_descriptive_statistics_table(data_path, algorithm_instances, document):
    document.add_paragraph()
    table = document.add_table(rows=len(algorithm_instances) + 1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Algorithm'
    header_cells[1].text = 'Mean'
    header_cells[2].text = 'Standard deviation'
    header_cells[3].text = 'Median'
    header_cells[4].text = 'Confidence interval'
    for i in range(len(algorithm_instances)):
        data = pd.read_csv(f'{data_path}/{algorithm_instances[i]}/Best so far.csv', header=None, index_col=None)
        mean, stdev, median, confidence_interval = get_descriptive_statistics(data.iloc[len(data) - 1])
        j = i + 1
        table.rows[j].cells[0].text = algorithm_instances[i]
        table.rows[j].cells[1].text = str(round(mean, 2))
        table.rows[j].cells[2].text = str(round(stdev, 2))
        table.rows[j].cells[3].text = str(round(median, 2))
        table.rows[j].cells[4].text = f'\u00B1 {round(confidence_interval, 2)}'


def _generate_report(experiment_number, algorithm_name, algorithm_instances, instances_indices):
    document = Document()
    for instance_idx in instances_indices:
        instance_number = instance_idx + 1
        plots_path = get_plots_path(experiment_number, instance_number, algorithm_name)
        document.add_heading(f'Instance {instance_number}', level=3)
        table = document.add_table(rows=len(algorithm_instances), cols=2)
        for i in range(len(algorithm_instances)):
            _add_image_to_cell(
                f'{plots_path}/{algorithm_instances[i]}/Best so far.png',
                table.rows[i].cells[0],
                2.98, 1.77
            )
            _add_image_to_cell(
                f'{plots_path}/{algorithm_instances[i]}/Current best.png',
                table.rows[i].cells[1],
                2.98, 1.77
            )

    for instance_idx in instances_indices:
        instance_number = instance_idx + 1
        _add_page_break(document)
        plots_path = get_plots_path(experiment_number, instance_number, algorithm_name)
        data_path = get_data_path(experiment_number, instance_number, algorithm_name)
        document.add_heading(f'Instance {instance_number} Final Results', level=3)
        _add_descriptive_statistics_table(data_path, algorithm_instances, document)
        document.add_paragraph()
        table = document.add_table(rows=2, cols=1)
        _add_image_to_cell(
            f'{plots_path}/Box Plots.png',
            table.rows[0].cells[0],
            6.11, 3.36,
        )
        _add_image_to_cell(
            f'{plots_path}/Confidence Intervals.png',
            table.rows[1].cells[0],
            6.11, 3.36,
        )

    document.save(f'experiments/Experiment {experiment_number}/{algorithm_name}.docx')


def generate_final_result_report(experiment_number, instances_indices):
    document = Document()
    for instance_idx in instances_indices:
        instance_number = instance_idx + 1
        _add_page_break(document)
        plots_path = get_plots_path(experiment_number, instance_number)
        document.add_heading(f'Instance {instance_number} Final Results', level=3)
        document.add_paragraph()
        table = document.add_table(rows=2, cols=1)
        _add_image_to_cell(
            f'{plots_path}/Box Plots.png',
            table.rows[0].cells[0],
            6.11, 3.9715,
        )
        _add_image_to_cell(
            f'{plots_path}/Confidence Intervals.png',
            table.rows[1].cells[0],
            6.11, 3.9715,
        )
    document.save(f'experiments/Experiment {experiment_number}/Final Results.docx')


def generate_report(experiment_number, algorithm_name_to_instances, instances_indices):
    for algorithm_name, algorithms_instances in algorithm_name_to_instances.items():
        _generate_report(experiment_number, algorithm_name, algorithms_instances, instances_indices)
    generate_final_result_report(experiment_number, instances_indices)
